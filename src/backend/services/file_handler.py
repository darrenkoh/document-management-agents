"""File handling operations for reading files."""
import logging
import hashlib
import base64
import io
import subprocess
import time
import re
from pathlib import Path
from typing import Optional, List
from pypdf import PdfReader
from docx import Document
from PIL import Image
try:
    from pillow_heif import register_heif_opener
    register_heif_opener()
except ImportError:
    pass
from openai import OpenAI

try:
    from openai import OpenAI
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False

# Try to import pdf2image for PDF to image conversion
try:
    from pdf2image import convert_from_path
    HAS_PDF2IMAGE = True
except ImportError:
    HAS_PDF2IMAGE = False

from src.backend.utils.retry import retry_on_llm_failure, RetryError

logger = logging.getLogger(__name__)


class FileHandler:
    """Handles file operations including text extraction and file moving."""

    def __init__(self, source_paths: List[str], ollama_endpoint: str = "http://localhost:11434",
                 ocr_model: str = "deepseek-ocr:3b", ocr_timeout: int = 60,
                 max_ocr_pages: int = 12, max_retries: int = 3, retry_base_delay: float = 1.0,
                 ocr_provider: str = "ollama", ocr_num_predict: int = 12000,
                 chandra_endpoint: str = "http://localhost:11435",
                 chandra_model: str = "chandra", chandra_timeout: int = 300,
                 chandra_max_tokens: int = 8192, chandra_max_retries: int = 3,
                 chandra_retry_base_delay: float = 1.0, chandra_frequency_penalty: float = 0.02,
                 chandra_detect_repeat_tokens: bool = True, hunyuan_endpoint: str = "http://localhost:11434",
                 hunyuan_model: str = "tencent/HunyuanOCR", hunyuan_timeout: int = 1800,
                 hunyuan_max_tokens: int = 16384, hunyuan_max_retries: int = 3,
                 hunyuan_retry_base_delay: float = 1.0, exclude_paths: Optional[List[str]] = None):
        """Initialize file handler.

        Args:
            source_paths: List of source directories to read files from
            ollama_endpoint: Ollama API endpoint for OCR
            ocr_model: OCR model name (default: deepseek-ocr:3b)
            ocr_timeout: Timeout for OCR operations in seconds
            max_ocr_pages: Maximum number of PDF pages to process with OCR (default: 12)
            max_retries: Maximum number of retry attempts for failed API calls
            retry_base_delay: Base delay in seconds between retry attempts
            ocr_provider: OCR provider ('ollama', 'chandra', or 'hunyuan')
            ocr_num_predict: Maximum number of tokens to predict for OCR (default: 12000)
            chandra_endpoint: Chandra vLLM API endpoint
            chandra_model: Chandra model name
            chandra_timeout: Timeout for Chandra OCR operations
            chandra_max_tokens: Maximum tokens for Chandra OCR
            chandra_max_retries: Maximum retries for Chandra OCR
            chandra_retry_base_delay: Base delay for Chandra OCR retries
            chandra_frequency_penalty: Frequency penalty to reduce repetition in generated text
            chandra_detect_repeat_tokens: Whether to detect and retry on repetitive OCR output
            hunyuan_endpoint: HunyuanOCR vLLM API endpoint
            hunyuan_model: HunyuanOCR model name
            hunyuan_timeout: Timeout for HunyuanOCR operations
            hunyuan_max_tokens: Maximum tokens for HunyuanOCR
            hunyuan_max_retries: Maximum retries for HunyuanOCR
            hunyuan_retry_base_delay: Base delay for HunyuanOCR retries
            exclude_paths: Optional list of paths to exclude from ingestion (absolute or relative)
        """
        self.source_paths = [Path(path) for path in source_paths]
        self.exclude_paths = [Path(p) for p in (exclude_paths or [])]
        self.ollama_endpoint = ollama_endpoint
        self.ocr_model = ocr_model
        self.ocr_timeout = ocr_timeout
        self.ocr_num_predict = ocr_num_predict
        self.max_ocr_pages = max_ocr_pages
        self.max_retries = max_retries
        self.retry_base_delay = retry_base_delay

        # Chandra OCR configuration
        self.ocr_provider = ocr_provider
        self.chandra_endpoint = chandra_endpoint
        self.chandra_model = chandra_model
        self.chandra_timeout = chandra_timeout
        self.chandra_max_tokens = chandra_max_tokens
        self.chandra_max_retries = chandra_max_retries
        self.chandra_retry_base_delay = chandra_retry_base_delay
        self.chandra_frequency_penalty = chandra_frequency_penalty
        self.chandra_detect_repeat_tokens = chandra_detect_repeat_tokens

        # HunyuanOCR configuration
        self.hunyuan_endpoint = hunyuan_endpoint
        self.hunyuan_model = hunyuan_model
        self.hunyuan_timeout = hunyuan_timeout
        self.hunyuan_max_tokens = hunyuan_max_tokens
        self.hunyuan_max_retries = hunyuan_max_retries
        self.hunyuan_retry_base_delay = hunyuan_retry_base_delay

        # Initialize OCR clients
        # Use OpenAI client for Ollama OCR (assuming Ollama is running with OpenAI compatibility mode)
        self.ollama_ocr_client = OpenAI(base_url=ollama_endpoint, api_key="dummy", timeout=ocr_timeout)
        self.chandra_ocr_client = None
        self.hunyuan_ocr_client = None
        if HAS_OPENAI:
            if ocr_provider == "chandra":
                self.chandra_ocr_client = OpenAI(
                    base_url=chandra_endpoint + "/v1",
                    api_key="dummy"  # vLLM doesn't require a real API key
                )
            elif ocr_provider == "hunyuan":
                self.hunyuan_ocr_client = OpenAI(
                    base_url=hunyuan_endpoint + "/v1",
                    api_key="dummy"  # vLLM doesn't require a real API key
                )

        self.ocr_available = self._check_ocr_availability()

        # Ensure source directories exist (skip if path is a file)
        for source_path in self.source_paths:
            if source_path.exists():
                if source_path.is_file():
                    # Path is a file, not a directory - skip mkdir
                    continue
                # Path exists and is a directory - no need to create
            else:
                # Path doesn't exist - create it as a directory
                source_path.mkdir(parents=True, exist_ok=True)

    def _get_ocr_provider_name(self) -> str:
        """Get the display name for the OCR provider.
        
        Returns:
            Display name for the OCR provider (e.g., "Chandra", "Hunyuan", or the actual model name for Ollama)
        """
        if self.ocr_provider == "chandra":
            return "Chandra"
        elif self.ocr_provider == "hunyuan":
            return "Hunyuan"
        else:
            # For Ollama provider, use the actual model name
            return self.ocr_model

    def _is_excluded(self, file_path: Path) -> bool:
        """Check if a file path should be excluded from ingestion."""
        if not self.exclude_paths:
            return False
        try:
            resolved = Path(file_path).resolve()
        except Exception:
            resolved = Path(file_path)
        for ex in self.exclude_paths:
            try:
                ex_resolved = ex.resolve()
            except Exception:
                ex_resolved = ex
            try:
                resolved.relative_to(ex_resolved)
                return True
            except Exception:
                continue
        return False

    def _call_ocr_generate(self, prompt: str, images: Optional[List[str]] = None, options: Optional[dict] = None) -> dict:
        """Make OCR LLM generate call with retry logic (deprecated - use provider-specific methods)."""
        if self.ocr_provider == "chandra":
            return self._call_chandra_ocr_generate(prompt, images, options)
        else:
            return self._call_ollama_ocr_generate(prompt, images, options)

    def _call_ollama_ocr_generate(self, prompt: str, images: Optional[List[str]] = None, options: Optional[dict] = None) -> dict:
        """Make OpenAI-compatible OCR LLM generate call with retry logic."""
        @retry_on_llm_failure(max_retries=self.max_retries,
                             base_delay=self.retry_base_delay,
                             exceptions=(Exception,))  # OCR can fail for various reasons
        def _generate():
            max_tokens = options.get('num_predict', 12000) if options else 12000
            temperature = options.get('temperature', 0.1) if options else 0.1

            # Build messages for vision API
            messages = [{"role": "user", "content": []}]

            # Add text prompt
            messages[0]["content"].append({"type": "text", "text": prompt})

            # Add images if provided
            if images:
                for image_path in images:
                    # Read image and encode as base64
                    import base64
                    with open(image_path, "rb") as image_file:
                        base64_image = base64.b64encode(image_file.read()).decode('utf-8')

                    messages[0]["content"].append({
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/jpeg;base64,{base64_image}"
                        }
                    })

            response = self.ollama_ocr_client.chat.completions.create(
                model=self.ocr_model,
                messages=messages,
                max_tokens=max_tokens,
                temperature=temperature
            )

            # Convert OpenAI response format to Ollama-like format for compatibility
            return {
                'response': response.choices[0].message.content,
                'done_reason': 'stop',  # Assume completion was successful
                'eval_count': None,     # Not available in OpenAI API
                'error': None
            }

        return _generate()

    def _call_chandra_ocr_generate(self, prompt: str, images: Optional[List[str]] = None, max_tokens: int = 8192) -> dict:
        """Make Chandra OCR LLM generate call with retry logic matching Chandra's implementation."""
        if not HAS_OPENAI or not self.chandra_ocr_client:
            raise Exception("OpenAI client not available for Chandra OCR")

        def _generate(temperature: float = 0.0, top_p: float = 0.1) -> dict:
            """Generate with specific temperature and top_p parameters."""
            content = []

            # Add image first (as per Chandra's implementation)
            if images:
                content.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:image/png;base64,{images[0]}"}
                })

            # Add text prompt after image
            content.append({"type": "text", "text": prompt})

            messages = [{"role": "user", "content": content}]

            response = self.chandra_ocr_client.chat.completions.create(
                model=self.chandra_model,
                messages=messages,
                max_tokens=max_tokens,
                temperature=temperature,
                top_p=top_p,
                frequency_penalty=self.chandra_frequency_penalty
            )

            # Convert OpenAI response format to Ollama format for compatibility
            return {
                'response': response.choices[0].message.content,
                'done': True
            }

        def _should_retry_chandra(result: dict, retries: int) -> bool:
            """Check if generation should be retried based on Chandra's logic."""
            raw_text = result.get('response', '')

            # Check for repeat tokens (Chandra's repeat detection) - if enabled
            if self.chandra_detect_repeat_tokens:
                has_repeat = self._detect_repeat_token(raw_text)
                if retries < self.chandra_max_retries and has_repeat:
                    logger.info(f"Detected repeat token in Chandra OCR output, retrying (attempt {retries + 1})")
                    return True

            # Check for other errors
            if retries < self.chandra_max_retries and not raw_text.strip():
                logger.info(f"Empty response from Chandra OCR, retrying (attempt {retries + 1})")
                return True

            return False

        # Main retry logic following Chandra's pattern
        result = _generate()  # Initial attempt with temperature=0, top_p=0.1
        retries = 0

        while _should_retry_chandra(result, retries):
            # Retry with higher temperature and top_p (as per Chandra's implementation)
            result = _generate(temperature=0.3, top_p=0.95)
            retries += 1
            if retries < self.chandra_max_retries:
                time.sleep(2 * (retries + 1))  # Exponential backoff

        return result

    def _call_hunyuan_ocr_generate(self, prompt: str, images: Optional[List[str]] = None, max_tokens: int = 16384) -> dict:
        """Make HunyuanOCR LLM generate call with retry logic."""
        if not HAS_OPENAI or not self.hunyuan_ocr_client:
            raise Exception("OpenAI client not available for HunyuanOCR")

        def _generate() -> dict:
            """Generate with HunyuanOCR Document Parsing prompt."""
            content = []

            # Add image first
            if images:
                content.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:image/png;base64,{images[0]}"}
                })

            # Add text prompt after image - using HunyuanOCR Document Parsing prompt
            content.append({"type": "text", "text": prompt})

            messages = [{"role": "user", "content": content}]

            response = self.hunyuan_ocr_client.chat.completions.create(
                model=self.hunyuan_model,
                messages=messages,
                max_tokens=max_tokens,
                temperature=0.0  # Deterministic output for OCR
            )

            # Convert OpenAI response format to Ollama format for compatibility
            return {
                'response': response.choices[0].message.content,
                'done': True
            }

        # Retry logic for HunyuanOCR
        @retry_on_llm_failure(max_retries=self.hunyuan_max_retries,
                             base_delay=self.hunyuan_retry_base_delay,
                             exceptions=(Exception,))
        def _generate_with_retry():
            return _generate()

        return _generate_with_retry()

    def _check_ocr_availability(self) -> bool:
        """Check if OCR is available and accessible."""
        if self.ocr_provider == "chandra":
            return self._check_chandra_availability()
        elif self.ocr_provider == "hunyuan":
            return self._check_hunyuan_availability()
        else:
            return self._check_ollama_availability()

    def _check_ollama_availability(self) -> bool:
        """Check if Ollama OCR is available and accessible."""
        try:
            # Try a simple test call to check availability
            self._call_ollama_ocr_generate(
                prompt="test",
                options={'num_predict': 1}  # Minimal response
            )
            return True
        except RetryError as e:
            logger.warning(f"Ollama OCR not available at {self.ollama_endpoint} after retries: {e.last_exception}")
            return False
        except Exception:
            logger.warning(f"Ollama OCR not available at {self.ollama_endpoint}")
            return False

    def _check_chandra_availability(self) -> bool:
        """Check if Chandra OCR is available and accessible."""
        if not HAS_OPENAI:
            logger.warning("OpenAI library not available for Chandra OCR")
            return False

        try:
            # Try a simple test call to check availability
            self._call_chandra_ocr_generate(
                prompt="test",
                max_tokens=1  # Minimal response
            )
            return True
        except RetryError as e:
            logger.warning(f"Chandra OCR not available at {self.chandra_endpoint} after retries: {e.last_exception}")
            return False
        except Exception as e:
            logger.warning(f"Chandra OCR not available at {self.chandra_endpoint}: {e}")
            return False

    def _check_hunyuan_availability(self) -> bool:
        """Check if HunyuanOCR is available and accessible."""
        if not HAS_OPENAI:
            logger.warning("OpenAI library not available for HunyuanOCR")
            return False

        try:
            # Try a simple test call to check availability
            self._call_hunyuan_ocr_generate(
                prompt="test",
                max_tokens=1  # Minimal response
            )
            return True
        except RetryError as e:
            logger.warning(f"HunyuanOCR not available at {self.hunyuan_endpoint} after retries: {e.last_exception}")
            return False
        except Exception as e:
            logger.warning(f"HunyuanOCR not available at {self.hunyuan_endpoint}: {e}")
            return False
    
    def _is_included(self, file_path: Path, allowed_extensions: Optional[List[str]] = None) -> bool:
        """Check if a file path should be included for processing.

        Args:
            file_path: Path to the file to check
            allowed_extensions: List of allowed file extensions (lowercase)

        Returns:
            True if the file should be included, False otherwise
        """
        if not allowed_extensions:
            return True  # If no extensions specified, include all files

        return file_path.suffix.lower() in allowed_extensions
    
    def get_files(self, extensions: Optional[List[str]] = None, recursive: bool = True) -> List[Path]:
        """Get list of files from source directories.

        Args:
            extensions: List of file extensions to include (e.g., ['.pdf', '.docx'])
                       If None, returns all files
            recursive: Whether to search recursively

        Returns:
            List of file paths that match the inclusion criteria
        """
        files = []

        for source_path in self.source_paths:
            if not source_path.exists():
                logger.warning(f"Source path does not exist: {source_path}")
                continue

            # Handle case where source_path is a file (not a directory)
            if source_path.is_file():
                if self._is_excluded(source_path):
                    continue
                # Check if file should be included
                if self._is_included(source_path, extensions):
                    files.append(source_path)
                else:
                    logger.info(f"Skipping file with disallowed extension: {source_path}")
                continue

            # source_path is a directory - search for files
            pattern = "**/*" if recursive else "*"

            for file_path in source_path.glob(pattern):
                if file_path.is_file():
                    if self._is_excluded(file_path):
                        continue
                    # Check if file should be included
                    if self._is_included(file_path, extensions):
                        files.append(file_path)
                    else:
                        logger.info(f"Skipping file with disallowed extension: {file_path}")
        
        return files
    
    def extract_text(self, file_path: Path) -> tuple[Optional[str], float, bool]:
        """Extract text content from a file.

        Args:
            file_path: Path to the file

        Returns:
            Tuple of (extracted text content or None if extraction fails, OCR duration in seconds, whether OCR was used)
        """
        # No try-except here to allow exceptions to propagate
        # Ensure file_path is a Path object
        if isinstance(file_path, str):
            file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix == '.pdf':
            text = self._extract_from_pdf(file_path)
            # Check if PDF text extraction returned meaningful content
            should_use_ocr = False
            if not text or not text.strip():
                should_use_ocr = True
            elif len(text.strip()) < 50:  # Threshold for "minimal" content
                should_use_ocr = True
            elif self._is_garbage_text(text):  # Check for garbage/random text
                should_use_ocr = True

            if should_use_ocr:
                if self.ocr_available:
                    ocr_provider_name = self._get_ocr_provider_name()
                    reason = "no content" if not text else "minimal content" if len(text.strip()) < 50 else "garbage text detected"
                    logger.info(f"PDF text extraction returned {reason}, trying {ocr_provider_name} for {file_path}")
                    ocr_result = self._extract_text_with_ocr(file_path)
                    if ocr_result:
                        ocr_text, ocr_duration = ocr_result
                        logger.info(f"{ocr_provider_name} successfully extracted text from {file_path}")
                        return (ocr_text, ocr_duration, True)
                    else:
                        logger.warning(f"{ocr_provider_name} failed, using extracted text (may be garbage)")
                        return (text, 0.0, False)
                else:
                    ocr_provider_name = self._get_ocr_provider_name()
                    logger.warning(f"{ocr_provider_name} not available, using potentially garbage text from {file_path}")
                    return (text, 0.0, False)
            return (text, 0.0, False)
        elif suffix == '.docx':
            text = self._extract_from_docx(file_path)
            # If DOCX text extraction returns minimal content, try OCR as fallback
            if not text or len(text.strip()) < 50:  # Threshold for "empty" content
                if self.ocr_available:
                    ocr_provider_name = self._get_ocr_provider_name()
                    logger.info(f"DOCX text extraction returned minimal content, trying {ocr_provider_name} for {file_path}")
                    ocr_result = self._extract_text_with_ocr(file_path)
                    if ocr_result:
                        ocr_text, ocr_duration = ocr_result
                        logger.info(f"{ocr_provider_name} successfully extracted text from DOCX {file_path}")
                        return (ocr_text, ocr_duration, True)
                    else:
                        return (text, 0.0, False)
                else:
                    ocr_provider_name = self._get_ocr_provider_name()
                    logger.warning(f"{ocr_provider_name} not available, skipping OCR fallback for {file_path}")
                    return (text, 0.0, False)
            return (text, 0.0, False)
        elif suffix == '.doc':
            return (self._extract_from_doc(file_path), 0.0, False)
        elif suffix == '.txt':
            return (self._extract_from_txt(file_path), 0.0, False)
        elif suffix in ['.png', '.jpg', '.jpeg', '.heic', '.gif', '.tiff', '.bmp']:
            if self.ocr_available:
                text, ocr_duration = self._extract_from_image(file_path)
                if text is not None:
                    return (text, ocr_duration, True)
                else:
                    logger.info(f"Skipping image {file_path} - no meaningful text content found")
                    return (None, ocr_duration, True)
            else:
                ocr_provider_name = self._get_ocr_provider_name()
                logger.warning(f"{ocr_provider_name} not available, skipping image {file_path}")
                return (None, 0.0, False)
        else:
            logger.warning(f"Unsupported file type: {suffix}")
            return (None, 0.0, False)

    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text_parts = []
        with open(file_path, 'rb') as f:
            pdf_reader = PdfReader(f)
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
        return '\n'.join(text_parts)
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                text_parts.append(paragraph.text)

            # Extract text from headers and footers
            for section in doc.sections:
                header = section.header
                if header:
                    for paragraph in header.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)

                footer = section.footer
                if footer:
                    for paragraph in footer.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)

            extracted_text = '\n'.join(text_parts)

            if not extracted_text.strip():
                logger.warning(f"No text content found in DOCX file {file_path}. File may be empty, corrupted, or contain only images.")
                return ""

            logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX file {file_path}")
            return extracted_text

        except Exception as e:
            logger.error(f"Failed to extract text from DOCX file {file_path}: {e}")
            return ""

    def _extract_from_doc(self, file_path: Path) -> str:
        """Extract text from DOC file using antiword."""
        try:
            # Try using antiword to extract text from .doc files
            result = subprocess.run(
                ['antiword', str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return result.stdout
            else:
                logger.warning(f"antiword failed for {file_path}: {result.stderr}")
                return ""
        except FileNotFoundError:
            logger.warning("antiword not installed. Install with: brew install antiword (macOS) or apt-get install antiword (Ubuntu)")
            return ""
        except subprocess.TimeoutExpired:
            logger.warning(f"antiword timed out for {file_path}")
            return ""
        except Exception as e:
            logger.error(f"Failed to extract text from DOC file {file_path}: {e}")
            return ""
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _extract_from_image(self, file_path: Path) -> tuple[Optional[str], float]:
        """Extract text from image using OCR.

        Returns:
            Tuple of (extracted text or None if no meaningful text found, OCR duration in seconds)
        """
        # No try-except here to allow exceptions to propagate
        image = Image.open(file_path)
        ocr_result = self._ocr_image(image)
        if ocr_result:
            text, duration = ocr_result
            # Check for meaningful content (not just whitespace or markdown boilerplate)
            stripped_text = text.strip() if text else ""
            if self._has_meaningful_content(stripped_text):
                ocr_provider_name = self._get_ocr_provider_name()
                logger.info(f"{ocr_provider_name} successfully extracted {len(stripped_text)} characters of meaningful content from {file_path}")
                return (text, duration)
            else:
                ocr_provider_name = self._get_ocr_provider_name()
                logger.info(f"{ocr_provider_name} found no meaningful text content in image {file_path} (likely a photo, chart, or empty image). Extracted: '{stripped_text[:100]}{'...' if len(stripped_text) > 100 else ''}'")
                return (None, duration)
        else:
            ocr_provider_name = self._get_ocr_provider_name()
            logger.warning(f"{ocr_provider_name} processing failed for {file_path}")
            return (None, 0.0)

    def _has_meaningful_content(self, text: str) -> bool:
        """Check if text contains meaningful content (not just boilerplate, formatting, or OCR metadata).

        Args:
            text: The text to check

        Returns:
            True if text appears to contain meaningful content, False otherwise
        """
        if not text:
            return False

        # Remove common markdown boilerplate
        text = text.replace('```markdown', '').replace('```', '').strip()

        # Remove DeepSeek OCR metadata tokens
        text = text.replace('<|ref|>', '').replace('<|/ref|>', '').replace('<|det|>', '').replace('<|/det|>', '').strip()

        # Remove coordinate data (patterns like [[x,y,w,h]])
        import re
        text = re.sub(r'\[\[\d+,\s*\d+,\s*\d+,\s*\d+\]\]', '', text).strip()

        # Remove HTML tags (common in some OCR outputs like Chandra/DeepSeek)
        # We replace with space to avoid merging words like <div>Word</div>
        clean_text = re.sub(r'<[^>]+>', ' ', text).strip()

        # Check for minimum length (at least 5 characters for meaningful content after stripping tags)
        if len(clean_text) < 5:
            return False

        # Check if it's just whitespace, punctuation, or common OCR artifacts
        if clean_text.replace(' ', '').replace('\n', '').replace('\t', '').replace('-', '').replace('_', '').replace('=', '').replace('*', '').replace('#', '') == '':
            return False

        # Check for common "no content" responses in the cleaned text
        no_content_indicators = [
            'no text found', 'no content', 'empty', 'blank',
            'no readable text', 'unable to extract', 'no data',
            'image', 'photo', 'picture', 'diagram', 'chart', 'graph'
        ]
        if any(indicator in clean_text.lower() for indicator in no_content_indicators):
            return False

        # Check if text contains actual readable words (not just symbols)
        # Filter for words that have at least one alphanumeric character and length >= 2
        words = [word for word in clean_text.split() if any(c.isalnum() for c in word) and len(word) >= 2]
        if len(words) < 2:  # Need at least 2 meaningful words
            return False

        return True

    def _is_garbage_text(self, text: str) -> bool:
        """Check if extracted text appears to be garbage/random characters rather than meaningful content.

        This is specifically designed to detect corrupted PDF text layers that contain
        random symbols instead of actual document content.

        Args:
            text: The text to check

        Returns:
            True if text appears to be garbage, False if it might be meaningful
        """
        if not text or not text.strip():
            return True

        # Remove whitespace for analysis
        cleaned_text = text.replace(' ', '').replace('\n', '').replace('\t', '').strip()

        # If too short after cleaning, it's likely empty
        if len(cleaned_text) < 10:
            return True

        # Calculate ratio of alphanumeric characters to total characters
        # Garbage text often has very low alphanumeric content
        alpha_numeric = sum(1 for c in cleaned_text if c.isalnum())
        alpha_numeric_ratio = alpha_numeric / len(cleaned_text)

        # If less than 30% alphanumeric characters, likely garbage
        if alpha_numeric_ratio < 0.3:
            return True

        # Check for excessive special character patterns common in corrupted PDFs
        # Patterns like repeated symbols, excessive punctuation, etc.
        special_chars = sum(1 for c in cleaned_text if not c.isalnum() and not c.isspace())
        special_ratio = special_chars / len(cleaned_text)

        # If more than 70% special characters, likely garbage
        if special_ratio > 0.7:
            return True

        # Check for repetitive patterns (common in corrupted text)
        # Look for sequences of 3+ identical characters
        for char in set(cleaned_text.lower()):
            if cleaned_text.lower().count(char * 3) > 0:
                return True

        # Check if text contains mostly symbols that don't form words
        words = text.split()
        if len(words) > 0:
            # Calculate what percentage of "words" are actually just symbols
            symbol_words = sum(1 for word in words if len(word) > 0 and all(not c.isalnum() for c in word))
            symbol_word_ratio = symbol_words / len(words)

            # If more than 50% of words are pure symbols, likely garbage
            if symbol_word_ratio > 0.5:
                return True

        # Check for common garbage patterns from corrupted PDFs
        garbage_patterns = [
            r'[~@#$%^&*+=|\\{}[\]:;"\'<>,.?/-]{5,}',  # 5+ consecutive special chars
            r'[a-zA-Z]{1,2}[~@#$%^&*+=|\\{}[\]:;"\'<>,.?/-]{3,}',  # short letters + many symbols
            r'[0-9]{1,3}[~@#$%^&*+=|\\{}[\]:;"\'<>,.?/-]{3,}',  # short numbers + many symbols
        ]

        for pattern in garbage_patterns:
            if re.search(pattern, text):
                return True

        return False

    def _detect_repeat_token(self, text: str, cut_from_end: int = None) -> bool:
        """Detect if text contains repeating tokens, following Chandra's implementation.

        Args:
            text: The text to check for repeats
            cut_from_end: If provided, only check the last N characters

        Returns:
            True if repeat tokens are detected
        """
        if not text:
            return False

        if cut_from_end:
            text = text[-cut_from_end:]

        # Look for sequences of 10+ identical characters
        for char in set(text):
            if text.count(char * 10) > 0:
                return True

        # Look for repeating patterns of 4+ characters
        for i in range(len(text) - 20):  # Check first 20 chars for patterns
            pattern = text[i:i+4]
            if len(pattern) >= 4 and text.count(pattern) > 3:  # Pattern repeats more than 3 times
                return True

        return False

    def _pdf_to_images(self, file_path: Path) -> List[Image.Image]:
        """Convert PDF pages to PIL Images for OCR processing.

        Args:
            file_path: Path to the PDF file

        Returns:
            List of PIL Images, one per page
        """
        try:
            if not HAS_PDF2IMAGE:
                logger.warning("pdf2image not available, cannot convert PDF to images for OCR")
                return []

            # Use pdf2image to convert PDF pages to images
            images = convert_from_path(file_path, dpi=200)  # 200 DPI for good OCR quality
            return images

        except Exception as e:
            logger.error(f"Failed to convert PDF to images: {e}")
            return []

    def _extract_text_with_ocr(self, file_path: Path) -> Optional[tuple[str, float]]:
        """Extract text from file using OCR as fallback.

        Args:
            file_path: Path to the file

        Returns:
            Tuple of (extracted text, duration) or None if OCR fails
        """
        # No try-except here to allow exceptions to propagate
        suffix = file_path.suffix.lower()
        ocr_provider_name = self._get_ocr_provider_name()

        if suffix == '.pdf':
            # Convert PDF to images first
            images = self._pdf_to_images(file_path)
            if not images:
                # If we can't extract images, try to render PDF pages as images
                # For now, return None - in production you might want to use pdf2image
                logger.warning(f"Could not extract images from PDF: {file_path}")
                return None

            # Limit the number of pages to process based on max_ocr_pages
            total_pages = len(images)
            if total_pages > self.max_ocr_pages:
                images = images[:self.max_ocr_pages]
                logger.info(f"PDF has {total_pages} pages, limiting OCR processing to first {self.max_ocr_pages} pages")

            # Process each image with OCR
            all_text = []
            total_ocr_duration = 0.0
            for i, image in enumerate(images):
                logger.info(f"Processing page {i+1}/{len(images)} with {ocr_provider_name}")
                ocr_result = self._ocr_image(image)
                if ocr_result:
                    text, duration = ocr_result
                    all_text.append(text)
                    total_ocr_duration += duration
                else:
                    logger.warning(f"Failed to extract text from page {i+1}")

            if all_text:
                return ('\n'.join(all_text), total_ocr_duration)
            else:
                return None

        elif suffix == '.docx':
            # For DOCX files, we need to convert to images first
            # This requires additional dependencies like docx2pdf and pdf2image
            logger.warning(f"DOCX OCR not yet implemented. Consider installing docx2pdf and pdf2image for DOCX OCR support.")
            return None

        elif suffix in ['.png', '.jpg', '.jpeg', '.gif', '.tiff', '.bmp']:
            # Direct image OCR
            image = Image.open(file_path)
            return self._ocr_image(image)  # Already returns (text, duration) or None
        else:
            return None

    def _ocr_image(self, image: Image.Image) -> Optional[tuple[str, float]]:
        """Perform OCR on a single image using the configured OCR provider.

        Args:
            image: PIL Image object

        Returns:
            Tuple of (extracted text, duration in seconds) or None if failed
        """
        if self.ocr_provider == "chandra":
            return self._ocr_image_with_chandra(image)
        elif self.ocr_provider == "hunyuan":
            return self._ocr_image_with_hunyuan(image)
        else:
            return self._ocr_image_with_deepseek(image)

    def _ocr_image_with_deepseek(self, image: Image.Image) -> Optional[tuple[str, float]]:
        """Perform OCR on a single image using Ollama OCR model.

        Args:
            image: PIL Image object

        Returns:
            Tuple of (extracted text, duration in seconds) or None if failed
        """
        # No try-except here to allow RetryError or Exception to propagate
        # Convert image to base64 string (not data URL)
        import io
        buffer = io.BytesIO()
        image.save(buffer, format='PNG')
        image_data = base64.b64encode(buffer.getvalue()).decode('utf-8')

        # Call Ollama OCR with timing and retry logic
        # Based on https://ollama.com/library/deepseek-ocr examples
        start_time = time.time()
        response = self._call_ollama_ocr_generate(
            prompt="<|grounding|>Convert the document to markdown.",
            images=[image_data],  # Base64 image data
            options={
                'temperature': 0.0,  # Deterministic output for OCR
                'num_predict': self.ocr_num_predict,  # Configurable token limit
            }
        )
        ocr_duration = time.time() - start_time

        # Check response field first
        extracted_text = response.get('response', '').strip()
        
        # For reasoning models, check thinking field as fallback if response is empty
        # Some reasoning models output to thinking field when they hit token limits
        if not extracted_text:
            thinking = response.get('thinking', '').strip()
            if thinking:
                # Check if thinking contains markdown patterns (final output)
                has_markdown = any(marker in thinking for marker in ['# ', '## ', '- ', '* ', '| ', '```'])
                
                # Also check the last portion of thinking for final output
                # Reasoning models sometimes put the answer at the end after reasoning
                last_portion = thinking[-2000:] if len(thinking) > 2000 else thinking
                has_final_output = any(marker in last_portion for marker in ['# ', '## ', '- ', '* ', '| ', '```'])
                
                if has_markdown or has_final_output:
                    # Try to extract just the markdown portion if possible
                    # Look for the start of markdown (first header or list)
                    markdown_start = -1
                    for marker in ['# ', '## ', '- ', '* ', '| ', '```']:
                        idx = thinking.find(marker)
                        if idx != -1 and (markdown_start == -1 or idx < markdown_start):
                            markdown_start = idx
                    
                    if markdown_start > 0:
                        # Extract from markdown start to end
                        extracted_text = thinking[markdown_start:].strip()
                        logger.info(f"{self._get_ocr_provider_name()} response was empty but found markdown in thinking field (starting at position {markdown_start}), using markdown portion")
                    else:
                        # Use entire thinking if we can't find a clear start
                        extracted_text = thinking
                        logger.info(f"{self._get_ocr_provider_name()} response was empty but found markdown in thinking field, using thinking content")
                else:
                    # If thinking is just reasoning without output, log it but don't use it
                    logger.warning(f"{self._get_ocr_provider_name()} response empty and thinking field contains only reasoning, not final output")
        
        ocr_provider_name = self._get_ocr_provider_name()
        if extracted_text:
            logger.info(f"{ocr_provider_name} extracted text: {extracted_text[:100]}...")
            return (extracted_text, ocr_duration)
        else:
            # Check if it was cut off due to token limit
            done_reason = response.get('done_reason', '')
            if done_reason == 'length':
                logger.warning(f"{ocr_provider_name} hit token limit (num_predict={self.ocr_num_predict}). Consider increasing num_predict in config.yaml")
            logger.info(f"{ocr_provider_name} returned empty response. Full response: {response}")
            return None

    def _ocr_image_with_chandra(self, image: Image.Image) -> Optional[tuple[str, float]]:
        """Perform OCR on a single image using Chandra-OCR.

        Args:
            image: PIL Image object

        Returns:
            Tuple of (extracted text, duration in seconds) or None if failed
        """
        # No try-except here to allow RetryError or Exception to propagate
        # Convert image to base64 string
        import io
        buffer = io.BytesIO()
        image.save(buffer, format='PNG')
        image_data = base64.b64encode(buffer.getvalue()).decode('utf-8')

        # Call Chandra-OCR with timing and retry logic
        # Based on Chandra documentation, it uses OpenAI-compatible API
        start_time = time.time()
        response = self._call_chandra_ocr_generate(
            prompt="Convert the document to markdown.",
            images=[image_data],  # Base64 image data
            max_tokens=self.chandra_max_tokens
        )
        ocr_duration = time.time() - start_time

        extracted_text = response.get('response', '').strip()
        if extracted_text:
            logger.info(f"Chandra-OCR extracted text: {extracted_text[:100]}...")
            return (extracted_text, ocr_duration)
        else:
            logger.info(f"Chandra-OCR returned empty response. Full response: {response}")
            return None

    def _ocr_image_with_hunyuan(self, image: Image.Image) -> Optional[tuple[str, float]]:
        """Perform OCR on a single image using HunyuanOCR.

        Args:
            image: PIL Image object

        Returns:
            Tuple of (extracted text, duration in seconds) or None if failed
        """
        # No try-except here to allow RetryError or Exception to propagate
        # Convert image to base64 string
        import io
        buffer = io.BytesIO()
        image.save(buffer, format='PNG')
        image_data = base64.b64encode(buffer.getvalue()).decode('utf-8')

        # Call HunyuanOCR with timing and retry logic
        # Using HunyuanOCR Document Parsing prompt
        start_time = time.time()
        response = self._call_hunyuan_ocr_generate(
            prompt="Extract all information from the main body of the document image and represent it in markdown format, ignoring headers and footers. Tables should be expressed in HTML format, formulas in the document should be represented using LaTeX format, and the parsing should be organized according to the reading order.",
            images=[image_data],  # Base64 image data
            max_tokens=self.hunyuan_max_tokens
        )
        ocr_duration = time.time() - start_time

        extracted_text = response.get('response', '').strip()
        if extracted_text:
            logger.info(f"HunyuanOCR extracted text: {extracted_text[:100]}...")
            return (extracted_text, ocr_duration)
        else:
            logger.info(f"HunyuanOCR returned empty response. Full response: {response}")
            return None

    def generate_file_hash(self, file_path: Path) -> Optional[tuple[str, float]]:
        """Generate SHA-256 hash of the file content.

        Args:
            file_path: Path to the file

        Returns:
            Tuple of (SHA-256 hash as hexadecimal string, duration in seconds) or None if failed
        """
        import time
        try:
            start_time = time.time()

            # Read the entire file in binary mode
            with open(file_path, 'rb') as f:
                file_content = f.read()

            # Generate SHA-256 hash
            hash_obj = hashlib.sha256(file_content)
            hash_value = hash_obj.hexdigest()

            duration = time.time() - start_time
            return (hash_value, duration)

        except Exception as e:
            logger.error(f"Error generating hash for {file_path}: {e}")
            return None
    

